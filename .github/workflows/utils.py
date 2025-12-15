# utils.py (Version Finale avec gestion ET/OU et debug)
import streamlit as st
import pandas as pd
import uuid
import firebase_admin
from firebase_admin import credentials, firestore
from datetime import datetime
import numpy as np
import zipfile
import io
import urllib.parse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL

# --- CONSTANTES ---
PROJECT_RENAME_MAP = {
    'Intitulé': 'Intitulé',
    'Fournisseur Bornes AC [Bornes]': 'Fournisseur Bornes AC',
    'Fournisseur Bornes DC [Bornes]': 'Fournisseur Bornes DC',
    'L [Plan de Déploiement]': 'PDC Lent',
    'R [Plan de Déploiement]': 'PDC Rapide',
    'UR [Plan de Déploiement]': 'PDC Ultra-rapide',
    'Pré L [Plan de Déploiement]': 'PDC L pré-équipés',
    'Pré R [Plan de Déploiement]': 'PDC R pré-équipés',
    'Pré UR [Plan de Déploiement]': 'PDC UR pré-équipés',
}

DISPLAY_GROUPS = [
    ['Intitulé', 'Fournisseur Bornes AC [Bornes]', 'Fournisseur Bornes DC [Bornes]'],
    ['L [Plan de Déploiement]', 'R [Plan de Déploiement]', 'UR [Plan de Déploiement]'],
    ['Pré L [Plan de Déploiement]', 'Pré R [Plan de Déploiement]', 'Pré UR [Plan de Déploiement]'],
]

SECTION_PHOTO_RULES = {
    "Bornes DC": ['R [Plan de Déploiement]', 'UR [Plan de Déploiement]'],
    "Bornes AC": ['L [Plan de Déploiement]'],
    # Ajoutez ici d'autres sections nécessitant des calculs de photos si nécessaire
}

COMMENT_ID = 100
COMMENT_QUESTION = "Veuillez préciser pourquoi le nombre de photo partagé ne correspond pas au minimum attendu"

# --- INITIALISATION FIREBASE ---
def initialize_firebase():
    """
    Initialise Firebase si ce n'est pas déjà fait. 
    Lit les secrets Streamlit au niveau racine (firebase_...).
    """
    if not firebase_admin._apps:
        try:
            # Lecture des secrets sans la section [firebase]
            cred_dict = {
                "type": st.secrets["firebase_type"],
                "project_id": st.secrets["firebase_project_id"],
                "private_key_id": st.secrets["firebase_private_key_id"],
                "private_key": st.secrets["firebase_private_key"].replace('\\n', '\n'),
                "client_email": st.secrets["firebase_client_email"],
                "client_id": st.secrets["firebase_client_id"],
                "auth_uri": st.secrets["firebase_auth_uri"],
                "token_uri": st.secrets["firebase_token_uri"],
                "auth_provider_x509_cert_url": st.secrets["firebase_auth_provider_x509_cert_url"],
                "client_x509_cert_url": st.secrets["firebase_client_x509_cert_url"],
                "universe_domain": st.secrets["firebase_universe_domain"],
            }
            
            project_id = cred_dict["project_id"]
            cred = credentials.Certificate(cred_dict)
            firebase_admin.initialize_app(cred, {'projectId': project_id})
        except Exception as e:
            st.error(f"Erreur de connexion Firebase : {e}")
            st.stop() 
            
    return firestore.client()

# Initialisation unique du client DB pour ce module
db = initialize_firebase()

# --- CHARGEMENT DONNÉES (mise en cache) ---
@st.cache_data(ttl=3600)
def load_form_structure_from_firestore():
    """Charge la structure du formulaire depuis la collection 'formsquestions'."""
    try:
        docs = db.collection('formsquestions').order_by('id').get()
        data = [doc.to_dict() for doc in docs]
        if not data: return None
        df = pd.DataFrame(data)
        df.columns = df.columns.str.strip()
        
        # Normalisation des noms de colonnes (résout les problèmes de typo)
        rename_map = {'Conditon value': 'Condition value', 'condition value': 'Condition value', 'Condition Value': 'Condition value', 'Condition': 'Condition value', 'Conditon on': 'Condition on', 'condition on': 'Condition on'}
        actual_rename = {k: v for k, v in rename_map.items() if k in df.columns}
        df = df.rename(columns=actual_rename)
        
        expected_cols = ['options', 'Description', 'Condition value', 'Condition on', 'section', 'id', 'question', 'type', 'obligatoire']
        for col in expected_cols:
            if col not in df.columns: df[col] = np.nan 
        
        # Nettoyage et typage
        df['options'] = df['options'].fillna('')
        df['Description'] = df['Description'].fillna('')
        df['Condition value'] = df['Condition value'].fillna('')
        df['Condition on'] = pd.to_numeric(df['Condition on'], errors='coerce').fillna(0).astype(int)
        
        for col in df.select_dtypes(include=['object']).columns:
            df[col] = df[col].astype(str).str.strip()
            try:
                # Tentative d'encodage/décodage pour gérer les caractères spéciaux (si nécessaire)
                df[col] = df[col].apply(lambda x: x.encode('utf-8', 'ignore').decode('utf-8', 'ignore'))
            except Exception: pass 
        return df
    except Exception as e:
        st.error(f"Erreur lors du chargement de la structure du formulaire: {e}")
        return None

@st.cache_data(ttl=3600)
def load_site_data_from_firestore():
    """Charge les données des sites depuis la collection 'Sites'."""
    try:
        docs = db.collection('Sites').get()
        data = [doc.to_dict() for doc in docs]
        if not data: return None
        df_site = pd.DataFrame(data)
        df_site.columns = df_site.columns.str.strip()
        return df_site
    except Exception as e:
        st.error(f"Erreur lors du chargement des données des sites: {e}")
        return None

# --- LOGIQUE MÉTIER ---

def get_expected_photo_count(section_name, project_data):
    """Calcule le nombre de photos attendues pour une section donnée."""
    if section_name.strip() not in SECTION_PHOTO_RULES:
        return None, None 

    columns = SECTION_PHOTO_RULES[section_name.strip()]
    total_expected = 0
    details = []

    for col in columns:
        val = project_data.get(col, 0)
        try:
            if pd.isna(val) or val == "":
                num = 0
            else:
                # Convertit la valeur en entier (gère les formats décimaux avec virgule/point)
                num = int(float(str(val).replace(',', '.'))) 
        except Exception:
            num = 0
        
        total_expected += num
        short_name = PROJECT_RENAME_MAP.get(col, col) 
        details.append(f"{num} {short_name}")

    detail_str = " + ".join(details)
    return total_expected, detail_str

def evaluate_single_condition(condition_str, all_answers):
    """
    Évalue une condition unitaire de type 'ID = Valeur'.
    Retourne True si la condition est respectée, False sinon.
    """
    if "=" not in condition_str:
        return True # Par défaut, si pas de format strict, on affiche
        
    try:
        target_id_str, expected_value_raw = condition_str.split('=', 1)
        target_id = int(target_id_str.strip())
        
        # Nettoie les guillemets/apostrophes pour la comparaison
        expected_value = expected_value_raw.strip().strip('"').strip("'")
        
        user_answer = all_answers.get(target_id)
        
        if user_answer is not None:
            # Comparaison insensible à la casse et aux espaces
            return str(user_answer).strip().lower() == str(expected_value).strip().lower()
        else:
            return False
    except Exception:
        return True # En cas d'erreur de parsing, on affiche par sécurité

def check_condition(row, current_answers, collected_data):
    """
    Vérifie si une question doit être affichée en fonction des réponses précédentes.
    Gère les opérateurs 'ET' et 'OU'.
    """
    try:
        # Si 'Condition on' n'est pas 1, la condition est inactive -> on affiche
        if int(row.get('Condition on', 0)) != 1: return True
    except (ValueError, TypeError): return True

    # Consolidation de toutes les réponses (passées et actuelles)
    all_past_answers = {}
    for phase_data in collected_data: 
        all_past_answers.update(phase_data['answers'])
    combined_answers = {**all_past_answers, **current_answers}
    
    # Nettoyage de la condition brute pour le parsing
    condition_raw = str(row.get('Condition value', '')).strip().strip('"').strip("'")
    if not condition_raw: return True

    # 1. Découpage par "OU" (Si un des blocs est Vrai, tout est Vrai)
    or_blocks = condition_raw.split(' OU ')
    
    for block in or_blocks:
        # 2. Découpage par "ET" (Dans un bloc, TOUT doit être Vrai)
        and_conditions = block.split(' ET ')
        block_is_valid = True
        
        for atom in and_conditions:
            # Utilisation de la fonction evaluate_single_condition
            if not evaluate_single_condition(atom, combined_answers):
                block_is_valid = False
                break # Une condition du ET est fausse, le bloc est faux
        
        if block_is_valid:
            return True # Un bloc entier est valide, donc la condition est remplie (grâce au OU)

    # Si aucun bloc n'a retourné True
    return False


def validate_section(df_questions, section_name, answers, collected_data, project_data):
    """Valide les réponses d'une section, y compris le compte de photos si applicable."""
    missing = []
    section_rows = df_questions[df_questions['section'] == section_name]
    
    comment_val = answers.get(COMMENT_ID)
    has_justification = comment_val is not None and str(comment_val).strip() != ""
    
    # 1. Calcul du nombre de photos attendu
    expected_total_base, detail_str = get_expected_photo_count(section_name.strip(), project_data)
    expected_total = expected_total_base
    
    photo_question_count = sum(
        1 for _, row in section_rows.iterrows()
        if str(row.get('type', '')).strip().lower() == 'photo' and check_condition(row, answers, collected_data)
    )
    
    if expected_total is not None and expected_total > 0:
        expected_total = expected_total_base * photo_question_count
        detail_str = (
            f"{detail_str} | Questions photo visibles: {photo_question_count} "
            f"-> Total ajusté: {expected_total}"
        )

    # 2. Compte des photos soumises
    current_photo_count = 0
    photo_questions_found = False
    
    for _, row in section_rows.iterrows():
        q_type = str(row['type']).strip().lower()
        if q_type == 'photo' and check_condition(row, answers, collected_data):
            photo_questions_found = True
            q_id = int(row['id'])
            val = answers.get(q_id)
            if isinstance(val, list):
                current_photo_count += len(val)

    # 3. Vérification des champs obligatoires (hors photos)
    for _, row in section_rows.iterrows():
        q_id = int(row['id'])
        if q_id == COMMENT_ID: continue
        if not check_condition(row, answers, collected_data): continue
        
        is_mandatory = str(row['obligatoire']).strip().lower() == 'oui'
        q_type = str(row['type']).strip().lower()
        val = answers.get(q_id)
        
        if is_mandatory and q_type != 'photo':
            if isinstance(val, list):
                if not val: missing.append(f"Question {q_id} : {row['question']} (fichier(s) manquant(s))")
            elif val is None or val == "" or (isinstance(val, (int, float)) and val == 0):
                missing.append(f"Question {q_id} : {row['question']}")

    # 4. Vérification de l'écart de photos et du commentaire
    is_photo_count_incorrect = False
    if expected_total is not None and expected_total > 0:
        if photo_questions_found and current_photo_count != expected_total:
            is_photo_count_incorrect = True
            error_message = (
                f"⚠️ **Écart de Photos pour '{str(section_name)}'**.\n"
                f"Attendu : **{str(expected_total)}** (calculé : {str(detail_str)}).\n"
                f"Reçu : **{str(current_photo_count)}**.\n"
            )
            if not has_justification:
                missing.append(
                    f"**Commentaire (ID {COMMENT_ID}) :** {COMMENT_QUESTION} "
                    f"(requis en raison de l'écart de photo : Attendu {expected_total}, Reçu {current_photo_count}).\n\n"
                    f"{error_message}"
                )

    # Nettoyage : Si le compte est bon, on retire le commentaire au cas où il ait été rempli par erreur
    if not is_photo_count_incorrect and COMMENT_ID in answers:
        del answers[COMMENT_ID]

    return len(missing) == 0, missing

# --- SAUVEGARDE ET EXPORTS (Fonctions inchangées) ---
def save_form_data(collected_data, project_data, submission_id, start_time):
    # ... (inchangé)
    try:
        cleaned_data = []
        for phase in collected_data:
            clean_phase = {
                "phase_name": phase["phase_name"],
                "answers": {}
            }
            # Remplace les objets FileUpload par des noms de fichiers
            for k, v in phase["answers"].items():
                if isinstance(v, list) and v and hasattr(v[0], 'read'): 
                    file_names = ", ".join([f.name for f in v])
                    clean_phase["answers"][str(k)] = f"Fichiers (non stockés en DB): {file_names}"
                
                elif hasattr(v, 'read'): 
                     clean_phase["answers"][str(k)] = f"Fichier (non stocké en DB): {v.name}"
                else:
                    clean_phase["answers"][str(k)] = v
            
            cleaned_data.append(clean_phase)
        
        final_document = {
            "project_intitule": project_data.get('Intitulé', 'N/A'),
            "project_details": project_data,
            "submission_id": submission_id,
            "start_date": start_time,
            "submission_date": datetime.now(),
            "status": "Completed",
            "collected_phases": cleaned_data
        }
        
        doc_id_base = str(project_data.get('Intitulé', 'form')).replace(" ", "_").replace("/", "_")[:20]
        doc_id = f"{doc_id_base}_{datetime.now().strftime('%Y%m%d_%H%M')}_{submission_id[:6]}"
        
        db.collection('FormAnswers').document(doc_id).set(final_document)
        return True, doc_id 
    except Exception as e:
        return False, str(e)

def create_csv_export(collected_data, df_struct, project_name, submission_id, start_time):
    # ... (inchangé)
    rows = []
    start_time_str = start_time.strftime('%Y-%m-%d %H:%M:%S') if isinstance(start_time, datetime) else 'N/A'
    end_time_str = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    for item in collected_data:
        phase_name = item['phase_name']
        for q_id, val in item['answers'].items():
            
            if int(q_id) == COMMENT_ID:
                q_text = "Commentaire Écart Photo"
            else:
                q_row = df_struct[df_struct['id'] == int(q_id)]
                q_text = q_row.iloc[0]['question'] if not q_row.empty else f"Question ID {q_id}"
            
            if isinstance(val, list) and val and hasattr(val[0], 'name'):
                final_val = f"[Pièces jointes] {len(val)} fichiers: " + ", ".join([f.name for f in val])
            elif hasattr(val, 'name'):
                final_val = f"[Pièce jointe] {val.name}"
            else:
                final_val = str(val)
            
            rows.append({
                "ID Formulaire": submission_id,
                "Date Début": start_time_str,
                "Date Fin": end_time_str,
                "Projet": project_name,
                "Phase": phase_name,
                "ID": q_id,
                "Question": q_text,
                "Réponse": final_val
            })
            
    df_export = pd.DataFrame(rows)
    return df_export.to_csv(index=False, sep=';', encoding='utf-8-sig')

def create_zip_export(collected_data):
    # ... (inchangé)
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        files_added = 0
        for phase in collected_data:
            phase_name_clean = str(phase['phase_name']).replace("/", "_").replace(" ", "_")
            for q_id, answer in phase['answers'].items():
                if isinstance(answer, list) and answer and hasattr(answer[0], 'read'):
                    for idx, file_obj in enumerate(answer):
                        try:
                            file_obj.seek(0)
                            file_content = file_obj.read()
                            if file_content:
                                original_name = file_obj.name.split('/')[-1].split('\\')[-1]
                                filename = f"{phase_name_clean}_Q{q_id}_{idx+1}_{original_name}"
                                zip_file.writestr(filename, file_content)
                                files_added += 1
                            file_obj.seek(0)
                        except Exception as e:
                            pass 
                            
        info_txt = f"Export généré le {datetime.now()}\nNombre de fichiers : {files_added}"
        zip_file.writestr("info.txt", info_txt)
    
    zip_buffer.seek(0)
    return zip_buffer

def define_custom_styles(doc):
    # ... (inchangé)
    # Style de titre principal
    try: title_style = doc.styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
    except: title_style = doc.styles['Report Title']
    title_style.base_style = doc.styles['Heading 1']
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(20)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0x01, 0x38, 0x2D)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(20)

    # Style de sous-titre de section
    try: subtitle_style = doc.styles.add_style('Report Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    except: subtitle_style = doc.styles['Report Subtitle']
    subtitle_style.base_style = doc.styles['Heading 2']
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Arial'
    subtitle_font.size = Pt(14)
    subtitle_font.bold = True
    subtitle_font.color.rgb = RGBColor(0x00, 0x56, 0x47)
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_style.paragraph_format.space_after = Pt(10)

    # Style de texte normal
    try: text_style = doc.styles.add_style('Report Text', WD_STYLE_TYPE.PARAGRAPH)
    except: text_style = doc.styles['Report Text']
    text_style.base_style = doc.styles['Normal']
    text_font = text_style.font
    text_font.name = 'Calibri'
    text_font.size = Pt(11)
    text_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    text_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    text_style.paragraph_format.space_after = Pt(5)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

def create_word_report(collected_data, df_struct, project_data, start_time):
    # ... (inchangé)
    doc = Document()
    define_custom_styles(doc)
    
    # --- Page de garde/Titre ---
    doc.add_paragraph('Rapport d\'Audit Chantier', style='Report Title')
    doc.add_paragraph('Informations du Projet', style='Report Subtitle')
    
    # Tableau d'informations de base
    project_table = doc.add_table(rows=3, cols=2)
    project_table.style = 'Light Grid Accent 1'
    project_table.rows[0].cells[0].text = 'Intitulé'
    project_table.rows[0].cells[1].text = str(project_data.get('Intitulé', 'N/A'))
    
    start_time_str = start_time.strftime('%d/%m/%Y %H:%M') if start_time else datetime.now().strftime('%d/%m/%Y %H:%M')
    project_table.rows[1].cells[0].text = 'Date de début'
    project_table.rows[1].cells[1].text = start_time_str
    project_table.rows[2].cells[0].text = 'Date de fin'
    project_table.rows[2].cells[1].text = datetime.now().strftime('%d/%m/%Y %H:%M')
    
    for row in project_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = 'Report Text'
    doc.add_paragraph()
    
    # Détails du projet (Bornes)
    doc.add_paragraph('Détails du Projet', style='Report Subtitle')
    for group in DISPLAY_GROUPS:
        for field_key in group:
            renamed_key = PROJECT_RENAME_MAP.get(field_key, field_key)
            value = project_data.get(field_key, 'N/A')
            p = doc.add_paragraph(style='Report Text')
            p.add_run(f'{renamed_key}: ').bold = True
            p.add_run(str(value))
    
    doc.add_page_break()
    
    # --- Contenu par Phase ---
    for phase_idx, phase in enumerate(collected_data):
        phase_name = phase['phase_name']
        doc.add_paragraph(f'Phase: {phase_name}', style='Report Subtitle')
        
        for q_id, answer in phase['answers'].items():
            # Récupération du texte de la question
            if int(q_id) == COMMENT_ID:
                q_text = "Commentaire explicatif de l'écart photo par rapport au nombre attendu"
            else:
                if not df_struct.empty:
                    q_row = df_struct[df_struct['id'] == int(q_id)]
                    q_text = q_row.iloc[0]['question'] if not q_row.empty else f"Question ID {q_id}"
                else:
                    q_text = f"Question ID {q_id}"
            
            # 1. Traitement des réponses de type PHOTO
            is_photo_answer = False
            if isinstance(answer, list) and answer and hasattr(answer[0], 'read'):
                is_photo_answer = True
            elif hasattr(answer, 'read'):
                is_photo_answer = True
                
            if is_photo_answer:
                doc.add_paragraph(f'Q{q_id}: {q_text}', style='Report Subtitle')
                if isinstance(answer, list):
                     doc.add_paragraph(f'Nombre de photos: {len(answer)}', style='Report Text')
                     for idx, file_obj in enumerate(answer):
                         try:
                             file_obj.seek(0)
                             image_data = file_obj.read()
                             if image_data:
                                 image_stream = io.BytesIO(image_data)
                                 doc.add_picture(image_stream, width=Inches(5)) 
                                 caption = doc.add_paragraph(f'Photo {idx+1}: {file_obj.name}', style='Report Text')
                                 caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                 caption.runs[0].font.size = Pt(9)
                                 caption.runs[0].font.italic = True
                                 file_obj.seek(0)
                         except Exception as e:
                             doc.add_paragraph(f'[Erreur photo {idx+1}: {e}]', style='Report Text')
                doc.add_paragraph()

            # 2. Traitement des autres types de réponses (texte, nombre, select)
            else:
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light Grid Accent 1'
                q_cell = table.cell(0, 0)
                q_cell.text = f'Q{q_id}: {q_text}'
                q_cell.width = Inches(5.0)
                a_cell = table.cell(0, 1)
                a_cell.text = str(answer)
                a_cell.width = Inches(1.0)
                for cell in table.rows[0].cells:
                    cell.paragraphs[0].style = 'Report Text'
                    cell.paragraphs[0].paragraph_format.left_indent = None 
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
                q_cell.paragraphs[0].runs[0].bold = True
                doc.add_paragraph()
        
        if phase_idx < len(collected_data) - 1:
            doc.add_page_break()
    
    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    return word_buffer

# --- COMPOSANT UI (rendu de la question) ---
def render_question(row, answers, phase_name, key_suffix, loop_index, project_data):
    """
    Rendu d'une question dans Streamlit, avec affichage de la condition et debug.
    """
    q_id = int(row.get('id', 0))
    is_dynamic_comment = q_id == COMMENT_ID
    
    # LIGNE DE DÉBOGAGE : Affiche les valeurs de la condition lues par Python pour TOUTES les questions rendues.
    # Note : Si une question est masquée par check_condition, cette fonction n'est pas appelée.
    if not is_dynamic_comment:
        # On utilise .get(key, '') pour s'assurer qu'une chaîne vide ou 'N/A' est affichée si la valeur est absente.
        st.write(
            f"DEBUG Q{q_id}: Cond. ON='{row.get('Condition on', 'N/A')}', "
            f"Cond. VALUE='{row.get('Condition value', 'N/A')}'"
        )
    
    if is_dynamic_comment:
        q_text = COMMENT_QUESTION
        q_type = 'text' 
        q_desc = "Ce champ est obligatoire si le nombre de photos n'est pas conforme."
        q_mandatory = True 
        q_options = []
        condition_display = ""
    else:
        q_text = row['question']
        q_type = str(row['type']).strip().lower() 
        q_desc = row['Description']
        q_mandatory = str(row['obligatoire']).lower() == 'oui'
        q_options = str(row['options']).split(',') if row['options'] else []
        
        # --- LOGIQUE POUR L'AFFICHAGE DE LA CONDITION ---
        condition_value = str(row.get('Condition value', '')).strip()
        condition_on = int(row.get('Condition on', 0))
        
        condition_display = ""
        if condition_on == 1 and condition_value:
             # Nettoyer les guillemets/apostrophes pour l'affichage (si la DB les inclut)
            display_value = condition_value.strip().strip('"').strip("'") 
            
            if display_value:
                condition_display = (
                    f'<span style="font-size: 0.8em; color: #a0a0a0; font-weight: normal; margin-left: 10px;">'
                    f'[Condition: {display_value}]'
                    f'</span>'
                )
        # ------------------------------------------------------------
        
    q_text = str(q_text).strip()
    q_desc = str(q_desc).strip()
    
    # Intégration de la condition dans le label HTML
    label_html = (
        f"<strong>{q_id}. {q_text}</strong>" 
        + (' <span class="mandatory">*</span>' if q_mandatory else "")
        + condition_display
    )
    
    widget_key = f"q_{q_id}_{phase_name}_{key_suffix}_{loop_index}"
    current_val = answers.get(q_id)
    val = current_val

    st.markdown(f'<div class="question-card"><div>{label_html}</div>', unsafe_allow_html=True)
    if q_desc: st.markdown(f'<div class="description">⚠️ {q_desc}</div>', unsafe_allow_html=True)
    
    # Éléments de formulaire
    if q_type == 'text':
        default_val = current_val if current_val else ""
        if is_dynamic_comment:
             val = st.text_area("Justification de l'écart", value=default_val, key=widget_key, label_visibility="collapsed")
        else:
             val = st.text_input("Réponse", value=default_val, key=widget_key, label_visibility="collapsed")

    elif q_type == 'select':
        clean_opts = [opt.strip() for opt in q_options]
        if "" not in clean_opts: clean_opts.insert(0, "")
        idx = clean_opts.index(current_val) if current_val in clean_opts else 0
        val = st.selectbox("Sélection", clean_opts, index=idx, key=widget_key, label_visibility="collapsed")
    
    elif q_type == 'number':
        label = "Nombre (Entier)"
        try:
            default_val = int(float(current_val)) if current_val is not None and str(current_val).replace('.', '', 1).isdigit() else 0
        except:
            default_val = 0
            
        val = st.number_input(
            label, 
            value=default_val, 
            step=1,             
            format="%d",         
            key=widget_key, 
            label_visibility="collapsed"
        )
    
    elif q_type == 'photo':
        expected, details = get_expected_photo_count(phase_name.strip(), project_data)
        if expected is not None and expected > 0:
            st.info(f"📸 **Photos :** Il est attendu **{expected}** photos pour cette section (Base calculée : {details}).")
            st.divider()
        
        file_uploader_default = current_val if isinstance(current_val, list) else []

        val = st.file_uploader("Images", type=['png', 'jpg', 'jpeg'], accept_multiple_files=True, key=widget_key, label_visibility="collapsed")
        
        if val:
            file_names = ", ".join([f.name for f in val])
            st.success(f"Nombre d'images chargées : {len(val)} ({file_names})")
        elif file_uploader_default and isinstance(file_uploader_default, list):
            val = file_uploader_default
            names = ", ".join([getattr(f, 'name', 'Fichier') for f in val])
            st.info(f"Fichiers conservés : {len(val)} ({names})")
        
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Mise à jour des réponses dans le dictionnaire 'answers'
    if val is not None and (not is_dynamic_comment or str(val).strip() != ""): 
        if q_type == 'number':
             answers[q_id] = int(val) 
        else:
            answers[q_id] = val 
    elif current_val is not None and not is_dynamic_comment: 
        answers[q_id] = current_val 
    elif is_dynamic_comment and (val is None or str(val).strip() == ""):
        if q_id in answers: del answers[q_id]
